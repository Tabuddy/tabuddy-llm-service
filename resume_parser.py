"""Resume file parser – extracts raw text from PDF and DOCX uploads."""

from __future__ import annotations

import io
import logging
from typing import Any

from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)


async def extract_text(file: UploadFile) -> str:
    """Read an uploaded file and return its full plain text."""
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _parse_pdf(content)
    elif filename.endswith(".docx"):
        return _parse_docx(content)
    elif filename.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Upload a .pdf, .docx, or .txt resume.",
        )


def _parse_pdf(data: bytes) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    full = "\n".join(text_parts).strip()
    if not full:
        raise HTTPException(
            status_code=422, detail="Could not extract text from PDF.")
    return full


def extract_pdf_links(data: bytes) -> list[dict[str, Any]]:
    """Return URI link annotations from a PDF (one entry per annotation).

    Uses pdfplumber link annotations (``/URI`` in link actions). Same PDFs as
    :func:`_parse_pdf`; no extra dependencies.
    """
    import pdfplumber

    def _inter_area(
        ax0: float,
        atop: float,
        ax1: float,
        abottom: float,
        bx0: float,
        btop: float,
        bx1: float,
        bbottom: float,
    ) -> float:
        ix0 = max(ax0, bx0)
        itop = max(atop, btop)
        ix1 = min(ax1, bx1)
        ibottom = min(abottom, bbottom)
        iw = max(0.0, ix1 - ix0)
        ih = max(0.0, ibottom - itop)
        return iw * ih

    def _v_overlap_frac(
        ltop: float,
        lbottom: float,
        wtop: float,
        wbottom: float,
    ) -> float:
        ih = max(0.0, min(lbottom, wbottom) - max(ltop, wtop))
        denom = min(lbottom - ltop, wbottom - wtop)
        if denom <= 0:
            return 0.0
        return ih / denom

    def _anchor_words_for_link(
        lx0: float,
        ltop: float,
        lx1: float,
        lbottom: float,
        words: list[dict[str, Any]],
    ) -> str:
        """Words inside the link rect, in an expanded box, or just to the right (icon+label)."""
        link_w = max(lx1 - lx0, 1e-6)
        link_h = max(lbottom - ltop, 1e-6)
        link_area = link_w * link_h
        pad_l = max(2.0, link_w * 0.05)
        pad_r = max(24.0, min(140.0, link_w * 4.0))
        pad_y = max(2.5, link_h * 0.35)
        exp_x0 = lx0 - pad_l
        exp_x1 = lx1 + pad_r
        exp_top = ltop - pad_y
        exp_bottom = lbottom + pad_y

        picked: list[tuple[float, float, str]] = []
        seen: set[tuple[int, int, str]] = set()

        for w in words:
            if not all(key in w for key in ("x0", "top", "x1", "bottom", "text")):
                continue
            text = str(w["text"]).strip()
            if not text:
                continue
            wx0 = float(w["x0"])
            wtop = float(w["top"])
            wx1 = float(w["x1"])
            wbottom = float(w["bottom"])
            word_area = max((wx1 - wx0) * (wbottom - wtop), 1e-6)

            inter = _inter_area(lx0, ltop, lx1, lbottom, wx0, wtop, wx1, wbottom)
            inter_exp = _inter_area(
                exp_x0, exp_top, exp_x1, exp_bottom, wx0, wtop, wx1, wbottom,
            )
            v_frac = _v_overlap_frac(ltop, lbottom, wtop, wbottom)

            frac_word = inter / word_area
            frac_link = inter / link_area

            same_line = v_frac >= 0.35
            # Icon + label: label often sits just right (or just left) of the clickable rect
            to_the_right = same_line and (wx0 >= lx1 - 1.5) and (wx0 <= lx1 + 130.0)
            to_the_left = same_line and (wx1 <= lx0 + 1.5) and (wx1 >= lx0 - 130.0)
            cx = 0.5 * (wx0 + wx1)
            cy = 0.5 * (wtop + wbottom)
            center_in_exp = exp_x0 <= cx <= exp_x1 and exp_top <= cy <= exp_bottom

            if (
                frac_word >= 0.12
                or frac_link >= 0.18
                or (inter_exp > 0 and inter_exp >= 0.25 * word_area)
                or to_the_right
                or to_the_left
                or (center_in_exp and v_frac >= 0.25)
            ):
                key = (round(wx0 * 2), round(wtop * 2), text)
                if key in seen:
                    continue
                seen.add(key)
                picked.append((wtop, wx0, text))

        if not picked:
            return ""
        picked.sort(key=lambda t: (t[0], t[1]))
        return " ".join(t[2] for t in picked).strip()

    out: list[dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            words = page.extract_words() or []
            for h in page.hyperlinks:
                uri = h.get("uri")
                if not uri or not str(uri).strip():
                    continue
                item: dict[str, Any] = {
                    "uri": str(uri).strip(),
                    "page": int(h["page_number"]),
                }
                if h.get("title") is not None:
                    item["title"] = h["title"]
                for k in ("x0", "y0", "x1", "y1", "top", "bottom", "width", "height"):
                    if k in h and h[k] is not None:
                        item[k] = float(h[k])

                if all(k in h and h[k] is not None for k in ("x0", "top", "x1", "bottom")):
                    lx0 = float(h["x0"])
                    ltop = float(h["top"])
                    lx1 = float(h["x1"])
                    lbottom = float(h["bottom"])
                    anchor = _anchor_words_for_link(lx0, ltop, lx1, lbottom, words)
                    if anchor:
                        uri_s = str(uri).strip()
                        if uri_s.lower().startswith("mailto:") and "|" in anchor:
                            addr = uri_s.split(":", 1)[1].split("?")[0].strip().lower()
                            for chunk in anchor.split("|"):
                                c = chunk.strip()
                                if addr and addr in c.lower().replace(" ", ""):
                                    anchor = c
                                    break
                            else:
                                anchor = anchor.split("|", 1)[0].strip()
                        item["anchor_text"] = anchor
                out.append(item)
    return out


def _table_cell_text_parts(cell: Any) -> list[str]:
    """Paragraphs and nested tables inside one table cell."""
    out: list[str] = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)
    for nested in cell.tables:
        out.extend(_table_text_parts(nested))
    return out


def _table_text_parts(table: Any) -> list[str]:
    out: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            out.extend(_table_cell_text_parts(cell))
    return out


def _header_footer_text_parts(doc: Any) -> list[str]:
    out: list[str] = []
    try:
        for section in doc.sections:
            for part in (section.header, section.footer):
                if part is None:
                    continue
                for p in part.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        out.append(t)
                for tbl in part.tables:
                    out.extend(_table_text_parts(tbl))
    except Exception as e:
        logger.debug("DOCX header/footer extraction skipped: %s", e)
    return out


def extract_docx_plain_text(data: bytes) -> str:
    """Best-effort plain text from a DOCX (body paragraphs, tables, headers/footers).

    Does not raise; returns \"\" if the file cannot be read or has no extractable text.
    Use :func:`_parse_docx` when an empty result should be a 422 for strict APIs.
    """
    import docx

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        logger.warning("DOCX could not be opened: %s", e)
        return ""

    chunks: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)
    for table in doc.tables:
        chunks.extend(_table_text_parts(table))
    chunks.extend(_header_footer_text_parts(doc))

    return "\n".join(chunks).strip()


def _parse_docx(data: bytes) -> str:
    full = extract_docx_plain_text(data)
    if not full:
        raise HTTPException(
            status_code=422, detail="Could not extract text from DOCX."
        )
    return full
